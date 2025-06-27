from flask import Flask, render_template, request, redirect, url_for, send_file
from flask_sqlalchemy import SQLAlchemy
from docx import Document
from datetime import datetime
import os

app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///' + os.path.join(os.getcwd(), 'database.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)

class Colaborador(db.Model):
    __tablename__ = 'colaborador'
    id = db.Column(db.Integer, primary_key=True)
    nome = db.Column(db.String)
    base = db.Column(db.String)
    cargo = db.Column(db.String)
    patrimonio = db.Column(db.String)
    modelo = db.Column(db.String)
    imei = db.Column(db.String)
    numero = db.Column(db.String)

@app.route('/')
def index():
    colaboradores = Colaborador.query.order_by(Colaborador.nome).all()
    return render_template('index.html', colaboradores=colaboradores)

@app.route('/termo/<int:id>', methods=['GET', 'POST'])
def formulario_termo(id):
    colaborador = Colaborador.query.get_or_404(id)
    
    if request.method == 'POST':
        # Processar checklist
        respostas = {}
        for i in range(1, 25):
            opcao = request.form.get(f'opcao{i}')
            respostas[f'sim{i}'] = 'X' if opcao == 'sim' else ' '
            respostas[f'nao{i}'] = 'X' if opcao == 'nao' else ' '
        
        # Gerar documento
        tipo_termo = request.form.get('tipo_termo')
        modelo = "Entrega.docx" if tipo_termo == "entrega" else "Devolução.docx"
        
        doc = Document(modelo)
        marcadores = {
            '{nome}': colaborador.nome,
            '{modelo}': colaborador.modelo,
            '{marca}': extrair_marca(colaborador.modelo),
            '{patrimonio}': colaborador.patrimonio,
            '{imei}': colaborador.imei,
            '{sim}': '(X)' if tem_linha_telefonica(colaborador.numero) else '( )',
            '{não}': '( )' if tem_linha_telefonica(colaborador.numero) else '(X)',
            '{cargo}': colaborador.cargo,
            '{local}': colaborador.base,
            '{info_adicionais}': request.form.get('info_adicionais', ''),
            '{data}': datetime.now().strftime('%d/%m/%Y')
        }
        
        # Adiciona as respostas do checklist
        for i in range(1, 25):
            marcadores[f'{{sim{i}}}'] = f'({respostas[f"sim{i}"]})'
            marcadores[f'{{nao{i}}}'] = f'({respostas[f"nao{i}"]})'
        
        # Substituição no documento
        substituir_marcadores(doc, marcadores)
        
        nome_arquivo = f"Termo_{tipo_termo}_{colaborador.nome}.docx"
        doc.save(nome_arquivo)
        
        return send_file(nome_arquivo, as_attachment=True)
    
    return render_template('formulario_termo.html', colaborador=colaborador)

def substituir_marcadores(doc, marcadores):
    """Substitui todos os marcadores no documento, incluindo tabelas"""
    for p in doc.paragraphs:
        for key, value in marcadores.items():
            p.text = p.text.replace(key, value)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in marcadores.items():
                    cell.text = cell.text.replace(key, value)

def tem_linha_telefonica(numero):
    """Verifica se o colaborador tem linha telefônica"""
    return numero and numero.strip() not in ['', 'NÃO POSSUI', '-']

def extrair_marca(modelo):
    """Extrai a marca do modelo do celular"""
    if not modelo:
        return ''
    modelo = modelo.lower()
    if 'galaxy' in modelo: return 'Samsung'
    if 'moto' in modelo: return 'Motorola'
    if 'iphone' in modelo: return 'Apple'
    return 'Outra'

if __name__ == '__main__':
    app.run(debug=True)